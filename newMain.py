# Example usage
import pandas as pd
from pprint import pprint as pp
from docx import Document
from datetime import datetime
import re
class WordDocumentFiller:
    def __init__(self):
        self.data = pd.DataFrame()
        self.template = None    
        
    def load_from_excel(self, file_path):
        """Load participant data from an Excel file."""
        print(f"Loading data from {file_path}...")
        self.data = pd.read_excel(file_path)
        print(f"Data loaded from {file_path}")

    def load_template(self, path):
        self.template = Document(path)
        print(f"Template loaded from {path}")

    def fill_template(self, center_id, input_date=""):
        """Fill the Word document template with data for a specific center ID."""
        row = self.getRowbyCenterID(center_id)
        if row is None:
            return None
            
        # Handle empty/NaN values by converting to empty string
        def safe_str(value):
            if pd.isna(value) or value is None:
                return ""
            return str(value)
        """ 
        import re

text = "Your sample text here with phone number like Liu, Lealie-daughter-9175139188 or (912)-112-2112."

# Regular expression pattern to match different phone number formats
pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'

# Find all matches
phone_numbers = re.findall(pattern, text) 
        """
        pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        emergency_text = safe_str(row["Emergency"])
        match = re.search(pattern, emergency_text)
        phone_number = match.group() if match else ""

        placeholder_map = {
            "{FULL_NAME}": safe_str(row["First_Name"]) + " " + safe_str(row["Last_Name"]),
            "{CHINESE_NAME}": safe_str(row["Chinese_Name"]),
            "{DOB}": safe_str(row["DOB"]),
            "{ADDRESS}": safe_str(row["Address"]),
            "{LANGUAGE}": safe_str(row["Language"]),
            "{MEDICAID_ID}": safe_str(row["Medicaid"]),
            "{GENDER}": safe_str(row["Gender"]),
            "{PCP}": safe_str(row["PCP"]),
            "{NAME}": safe_str(row["Emergency"]),  # Emergency contact name
            "{CURRENT_DATE}": input_date if input_date else datetime.now().strftime("%m/%d/%Y"),
            "{COMPANY_ID}": safe_str(row["Member_ID"]),
            "{COMPANY}": safe_str(row["Health_Plan"]),
            "{PHONE}": safe_str(row["Home_Tel"])  if safe_str(row["Home_Tel"]) != "" else safe_str(row["Cell"]),
            "{MEDICARE_ID}": safe_str(row["Medicare"]),
            "{EMERGENCY_PHONE}":  phone_number  # Use .get() for optional fields
        }
        
        print("Filling template with the following data:")
        pp(placeholder_map)
        
        # DEBUG: Print all text in the document to see what placeholders exist
        print("\n🔍 DEBUG: All text found in document:")
        print("=" * 50)
        
        for i, paragraph in enumerate(self.template.paragraphs):
            if paragraph.text.strip():
                print(f"Paragraph {i}: '{paragraph.text}'")
        
        for table_idx, table in enumerate(self.template.tables):
            print(f"\nTable {table_idx}:")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        print(f"  Row {row_idx}, Cell {cell_idx}: '{cell.text}'")
        
        print("=" * 50)
        
        # Better approach: Replace text while preserving formatting
        replacements_made = self._replace_text_in_runs(placeholder_map)
        
        if replacements_made == 0:
            print("⚠️  WARNING: No placeholders were found and replaced!")
            print("Make sure your Word document contains placeholders like {FULL_NAME}, {CURRENT_DATE}, etc.")
        
        print(f"Template filled for Center_ID: {center_id}")
        return self.template

    def _replace_text_in_runs(self, placeholder_map):
        """Replace placeholders while preserving formatting by working with runs."""
        replacements_made = 0
        
        # Replace in paragraphs
        for paragraph in self.template.paragraphs:
            replacements_made += self._replace_paragraph_text(paragraph, placeholder_map)
        
        # Replace in tables
        for table in self.template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replacements_made += self._replace_paragraph_text(paragraph, placeholder_map)
        
        return replacements_made

    def _replace_paragraph_text(self, paragraph, placeholder_map):
        """Replace text in a paragraph while preserving formatting."""
        # Get the full text of the paragraph
        full_text = paragraph.text
        replacements_made = 0
        
        # Check if any placeholder exists in this paragraph
        has_placeholder = any(placeholder in full_text for placeholder in placeholder_map.keys())
        
        if has_placeholder:
            # Replace all placeholders in the full text
            new_text = full_text
            for placeholder, value in placeholder_map.items():
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, value)
                    print(f"✅ Replaced {placeholder} with '{value}' in: '{full_text[:50]}...'")
                    replacements_made += 1
            
            # Clear existing runs and add new text with preserved base formatting
            if paragraph.runs:
                # Keep the formatting of the first run
                first_run = paragraph.runs[0]
                # Clear all runs
                for run in paragraph.runs[::-1]:  # Reverse to avoid index issues
                    run.clear()
                # Set new text with original formatting
                paragraph.text = new_text
            else:
                paragraph.text = new_text
        
        return replacements_made

    def save_filled_document(self, output_path):
        """Save the filled document to a file."""
        if self.template:
            self.template.save(output_path)
            print(f"Document saved to: {output_path}")
        else:
            print("No template loaded to save.")

    def getRowbyCenterID(self, center_id):
        """Get a row of data by center ID."""
        row = self.data[self.data['Center_ID'] == center_id]
        if not row.empty:
            return row.iloc[0]
        else:
            print(f"No data found for Center_ID: {center_id}")
            return None

def main():
    # Initialize the filler
    filler = WordDocumentFiller()
    filler.load_from_excel('contact.xlsx')
    
    while True:
        center_id = input("Enter Center_ID to retrieve data (or 'exit' to quit): ")
        if center_id.lower() == 'exit' or center_id.lower() == "quit":
            break
        try:
            center_id = int(center_id)
            row_data = filler.getRowbyCenterID(center_id)
            if row_data is not None:
                filler.load_template('templateCopy.docx')
                filled_template = filler.fill_template(center_id, input_date="01/01/2021")
                
                if filled_template:
                    # Save the filled document
                    output_filename = f"filled_form_{center_id}.docx"
                    filler.save_filled_document(output_filename)
                    print(f"✅ Document created: {output_filename}")
                    
                    # Show what was actually replaced
                    print("\n📋 Summary of replacements made:")
                    for placeholder, value in [("{FULL_NAME}", f"{row_data['First_Name']} {row_data['Last_Name']}"),
                                              ("{CURRENT_DATE}", datetime.now().strftime("%m/%d/%Y"))]:
                        if value.strip():
                            print(f"  {placeholder} → {value}")
                    
        except ValueError:
            print("Please enter a valid Center_ID or 'exit' to quit.")
        except Exception as e:
            print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()