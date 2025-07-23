# Example usage
import pandas as pd
from pprint import pprint as pp
from docx import Document
from datetime import datetime

class WordDocumentFiller:
    def __init__(self):
        self.data = pd.DataFrame()
        self.template = None    
        
    def load_from_excel(self, file_path):
        """Load participant data from an Excel file."""
        print(f"Loading data from {file_path}...")
        self.data = pd.read_excel(file_path)
        print(f"Data loaded from {file_path}")

    def load_template(self, path):
        self.template = Document(path)
        print(f"Template loaded from {path}")

    def fill_template(self, center_id, input_date=""):
        """Fill the Word document template with data for a specific center ID."""
        row = self.getRowbyCenterID(center_id)
        if row is None:
            return None
            
        # Handle empty/NaN values by converting to empty string
        def safe_str(value):
            if pd.isna(value) or value is None:
                return ""
            return str(value)
        
        placeholder_map = {
            "{FULL_NAME}": safe_str(row["First_Name"]) + " " + safe_str(row["Last_Name"]),
            "{CHINESE_NAME}": safe_str(row["Chinese_Name"]),
            "{DOB}": safe_str(row["DOB"]),
            "{ADDRESS}": safe_str(row["Address"]),
            "{LANGUAGE}": safe_str(row["Language"]),
            "{MEDICAID_ID}": safe_str(row["Medicaid"]),
            "{GENDER}": safe_str(row["Gender"]),
            "{PCP}": safe_str(row["PCP"]),
            "{NAME}": safe_str(row["Emergency"]),  # Emergency contact name
            "{CURRENT_DATE}": input_date if input_date else datetime.now().strftime("%m/%d/%Y"),
            "{COMPANY_ID}": safe_str(row["Member_ID"]),
            "{COMPANY}": safe_str(row["Health_Plan"]),
            "{PHONE}": safe_str(row["Home_Tel"]),
            "{MEDICARE_ID}": safe_str(row["Medicare"]),
            "{EMERGENCY_PHONE}": safe_str(row.get("Emergency_Phone", ""))  # Use .get() for optional fields
        }
        
        print("Filling template with the following data:")
        pp(placeholder_map)
        
        # Replace placeholders in paragraphs
        for paragraph in self.template.paragraphs:
            for placeholder, value in placeholder_map.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
                    print(f"Replaced {placeholder} with {value}")
        
        # Replace placeholders in tables
        for table in self.template.tables:
            for row_table in table.rows:
                for cell in row_table.cells:
                    for placeholder, value in placeholder_map.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
                            print(f"Replaced {placeholder} with {value} in table")
        
        print(f"Template filled for Center_ID: {center_id}")
        return self.template

    def save_filled_document(self, output_path):
        """Save the filled document to a file."""
        if self.template:
            self.template.save(output_path)
            print(f"Document saved to: {output_path}")
        else:
            print("No template loaded to save.")

    def getRowbyCenterID(self, center_id):
        """Get a row of data by center ID."""
        row = self.data[self.data['Center_ID'] == center_id]
        if not row.empty:
            return row.iloc[0]
        else:
            print(f"No data found for Center_ID: {center_id}")
            return None

def main():
    # Initialize the filler
    filler = WordDocumentFiller()
    filler.load_from_excel('contact.xlsx')
    
    while True:
        center_id = input("Enter Center_ID to retrieve data (or 'exit' to quit): ")
        if center_id.lower() == 'exit' or center_id.lower() == "quit":
            break
        try:
            center_id = int(center_id)
            row_data = filler.getRowbyCenterID(center_id)
            if row_data is not None:
                filler.load_template('templateCopy.docx')
                filled_template = filler.fill_template(center_id, input_date="01/01/2021")
                
                if filled_template:
                    # Save the filled document
                    output_filename = f"filled_form_{center_id}.docx"
                    filler.save_filled_document(output_filename)
                    print(f"✅ Document created: {output_filename}")
                    
                    # Show what was actually replaced
                    print("\n📋 Summary of replacements made:")
                    for placeholder, value in [("{FULL_NAME}", f"{row_data['First_Name']} {row_data['Last_Name']}"),
                                              ("{CURRENT_DATE}", datetime.now().strftime("%m/%d/%Y"))]:
                        if value.strip():
                            print(f"  {placeholder} → {value}")
                    
        except ValueError:
            print("Please enter a valid Center_ID or 'exit' to quit.")
        except Exception as e:
            print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()