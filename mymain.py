# Example usage
import pandas as pd
from pprint import pprint as pp
from docx import Document
import re
import os 

class WordDocumentFiller:
    def __init__(self,center_id):
        self.data = pd.DataFrame
        self.template = None
        self.center_id = center_id
    
    def load_from_excel(self, file_path):
        print(f"Loading data from {file_path}...")
        self.data = pd.read_excel(file_path)
        print(f"Data loaded from {file_path}")
        
    def load_template(self, path):
        self.template = Document(path)
        print(f"Template loaded from {path}")
    def getRowByCenterId(self, center_id):
        """Get a row of data by center ID."""
        row = self.data[self.data['Center_ID'] == center_id]
        if not row.empty:
            return row.iloc[0]
        else:
            print(f"No data found for Center_ID: {center_id}")
            return None
        
    def fill_template(self, input_date=""):
        seenParas = set()
        dataRow  = self.getRowByCenterId(self.center_id)
        if dataRow is None: return None

        def safe_str(value):
            if pd.isna(value) or value is None:
                return ""
            return str(value)
        pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        emergency_text = safe_str(dataRow["Emergency"])
        match = re.search(pattern, emergency_text)
        phone_number = match.group() if match else ""
        emergencyName = safe_str(dataRow["Emergency"])
        emergencyName = re.sub(r"[\d()\-\s]+", "", emergencyName)
        placeholder_map = {
            "{FULL_NAME}": safe_str(dataRow["First_Name"]) + " " + safe_str(dataRow["Last_Name"]),
            "{CHINESE_NAME}": safe_str(dataRow["Chinese_Name"]),
            "{DOB}": safe_str(dataRow["DOB"]),
            "{ADDRESS}": safe_str(dataRow["Address"]),
            "{LANGUAGE}": safe_str(dataRow["Language"]),
            "{MEDICAID_ID}": safe_str(dataRow["Medicaid"]),
            "{GENDER}": safe_str(dataRow["Gender"]),
            "{PCP}": safe_str(dataRow["PCP"]),
            "{NAME}": emergencyName,  # Emergency contact name
            "{CURRENT_DATE}": input_date,
            "{COMPANY_ID}": safe_str(dataRow["Member_ID"]),
            "{COMPANY}": safe_str(dataRow["Health_Plan"]),
            "{PHONE}": safe_str(dataRow["Home_Tel"])  if safe_str(dataRow["Home_Tel"]) != "" else safe_str(dataRow["Cell"]),
            "{MEDICARE_ID}": safe_str(dataRow["Medicare"]),
            "{EMERGENCY_PHONE}":  phone_number  # Use .get() for optional fields
        }

        def replace_in_paragraph(paragraph):
            for run in paragraph.runs:
                # print(f"text: {run.text} END")
                # for key, val in placeholder_map.items():
                if run.text in placeholder_map:
                    print(f"✅ Replaced {run.text} with '{placeholder_map[run.text]}' ")
                    run.text = placeholder_map[run.text]
                print(f"current text: {run.text} is in?: {'{CURRENT_DATE}' in run.text}")
                
                if "CURRENT_DATE}" in run.text:
                    run.text = placeholder_map["{CURRENT_DATE}"]
# Replace placeholders in document-level paragraphs
        # for para in self.template.paragraphs:
        #     replace_in_paragraph(para)

# Replace placeholders in tables, safely handling merged cells
        for table in self.template.tables:
            row_count = len(table.rows)
            col_count = len(table.columns)
            # for r in range(row_count):
            #     for c in range(col_count):
            #         try:
            #             cell = table.cell(r, c)
            #             print(cell.text)
            #             for para in cell.paragraphs:
            #                 # if para.text != "{CURRENT_DATE}" and para.text in seenParas: continue
            #                 seenParas.add(para.text)
            #                 replace_in_paragraph(para)
            for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        # print(cell.text)
                        for para in cell.paragraphs:
                            if para.text != "{CURRENT_DATE}" and para.text in seenParas: continue
                            seenParas.add(para.text)
                            replace_in_paragraph(para)
                            print(f"para: {para.text} End")
                        # if cell.text.strip():
                        #     print(f"  Row {row_idx}, Cell {cell_idx}: '{cell.text}'")
                    # except IndexError:
                        # continue  # skip over malformed cell references



        # for table_idx, table in enumerate(self.template.tables):
        #         print(f"\nTable {table_idx}:")
        #         for row_idx, row in enumerate(table.rows):
        #             for cell_idx, cell in enumerate(row.cells):
        #                 if cell.text.strip():
        #                     print(f"  Row {row_idx}, Cell {cell_idx}: '{cell.text}'")


        # for para in self.template:
        #     print(para)
        # for section_idx, section in enumerate(self.template.sections):
        #     print(section)
        #     a = section.footer.paragraphs
        #     for para in a:
        #         print(para.text)
            # for footer in section.footer:
            #     print(footer)
        
        # print_matching_paragraphs(footer.paragraphs, f"Footer {section_idx}")
        self.save_filled_document(f"PAF-{input_date}.docx")
    def save_filled_document(self, output_path):
        """Save the filled document to a file."""
        if self.template:
            self.template.save(output_path)
            print(f"Document saved to: {output_path}")
        else:
            print("No template loaded to save.")

def main():
    filler = WordDocumentFiller(100)
    filler.load_from_excel("ScriptContacts.xlsx")
    # filler.load_template('templateCopy.docx')
    filler.load_template('templateCopy.docx')
    # while True:
    #     center_id = input("Enter Center id: ")
    #     date = input("enter date string: ")
    filler.fill_template("12-12-1222")

if __name__ == "__main__":
    main()