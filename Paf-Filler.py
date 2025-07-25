# Example usage
import pandas as pd
from pprint import pprint as pp
from docx import Document
import re
from datetime import datetime
import os 
from pathlib import Path

BASEFOLDER  = ""


class WordDocumentFiller:
    def __init__(self):
        self.data = pd.DataFrame
        self.template = None
        self.center_id = -9999
        self.date = "UNKNOWN_DATE"
    def set_date(self,date):
        self.date = date
    def set_center_id(self,center_id):
        self.center_id = center_id

    def load_from_excel(self, file_path):
        print(f"Loading data from {file_path}...")
        self.data = pd.read_excel(file_path)
        print(f"Data loaded from {file_path}")
        
    def load_template(self, path):
        self.template = Document(path)
        print(f"Template loaded from {path}")
    def getRowByCenterId(self, center_id):
        """Get a row of data by center ID."""
        row = self.data[self.data['Center_ID'] == center_id]
        if not row.empty:
            return row.iloc[0]
        else:
            print(f"No data found for Center_ID: {center_id}")
            return None
        
    def fill_template(self):
        seenParas = set()
        dataRow  = self.getRowByCenterId(self.center_id)
        
        if dataRow is None: return False

        def safe_str(value):
            if pd.isna(value) or value is None:
                return ""
            return str(value)
        pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        emergency_text = safe_str(dataRow["Emergency"])
        match = re.search(pattern, emergency_text)
        phone_number = match.group() if match else ""
        emergencyName = safe_str(dataRow["Emergency"])
        emergencyName = re.sub(r"[\d()\-\s]+", "", emergencyName)
        placeholder_map = {
            "{FULL_NAME}": safe_str(dataRow["First_Name"]) + " " + safe_str(dataRow["Last_Name"]),
            "{CHINESE_NAME}": safe_str(dataRow["Chinese_Name"]),
            "{DOB}": safe_str(dataRow["DOB"]),
            "{ADDRESS}": safe_str(dataRow["Address"]),
            "{LANGUAGE}": safe_str(dataRow["Language"]),
            "{MEDICAID_ID}": safe_str(dataRow["Medicaid"]),
            "{GENDER}": safe_str(dataRow["Gender"]),
            "{PCP}": safe_str(dataRow["PCP"]),
            "{NAME}": emergencyName,  # Emergency contact name
            "{CURRENT_DATE}": self.date,
            "{COMPANY_ID}": safe_str(dataRow["Member_ID"]),
            "{COMPANY}": safe_str(dataRow["Health_Plan"]),
            "{PHONE}": safe_str(dataRow["Home_Tel"])  if safe_str(dataRow["Home_Tel"]) != "" else safe_str(dataRow["Cell"]),
            "{MEDICARE_ID}": safe_str(dataRow["Medicare"]),
            "{EMERGENCY_PHONE}":  phone_number  # Use .get() for optional fields
        }

        def replace_in_paragraph(paragraph):
            for run in paragraph.runs:
                if run.text in placeholder_map:
                    print(f"✅ Replaced {run.text} with '{placeholder_map[run.text]}' ")
                    run.text = placeholder_map[run.text]
                if "CURRENT_DATE}" in run.text:
                    run.text = placeholder_map["{CURRENT_DATE}"]
                    print(f"✅ Replaced {run.text} with '{placeholder_map[run.text]}' ")

# Replace placeholders in tables, safely handling merged cells
        for table in self.template.tables:
            for _, row in enumerate(table.rows):
                    for _, cell in enumerate(row.cells):
                        for para in cell.paragraphs:
                            if para.text != "{CURRENT_DATE}" and para.text in seenParas: continue
                            seenParas.add(para.text)
                            replace_in_paragraph(para)
     
        return True
    def save_filled_document(self, output_path):
        """Save the filled document to a file."""
        if self.template:
            self.template.save(output_path)
            print(f"Document saved to: {output_path}")
        else:
            print("No template loaded to save.")


def isValidDate(date_input="Enter date (e.g., 1/2/2025 or 01-02-2025): "):
    formats = ["%m/%d/%Y", "%m-%d-%Y"]
    for fmt in formats:
        try:
            _ =  datetime.strptime(date_input, fmt).strftime("%m/%d/%Y")
            return True
        except:
            pass
           
    return False

def main():
    filler = WordDocumentFiller()
    filler.load_from_excel("ScriptContacts.xlsx")
    
    filler.load_template('templateCopy.docx')
    while True:
        center_id = input("Enter Center id: ")
        if len(center_id) != 0 and center_id[0] == 'q' or  center_id[0] == 'e':
            print("exiting program")
            break
        try:
            center_id = int(center_id)
        except ValueError as e:
            print(f"error: {e}")
            continue
            
        date = input("Enter date string: ")
        if not isValidDate(date):
            print("Please enter valid date. {date} is invalid" )
            continue
        filler.set_center_id(center_id)
        filler.set_date(date)
        home = Path.home()
        dirpath = home / "OneDrive" / "Desktop" / "Alden" / "member_profiles" / f"{center_id}"
        Path(dirpath).mkdir(exist_ok=True)
        outpath = dirpath / f"PAF-{filler.date}.docx"
        
        isFilled = filler.fill_template()
        if isFilled:
            filler.save_filled_document(outpath)
            os.startfile(outpath)
        
if __name__ == "__main__":
    main()